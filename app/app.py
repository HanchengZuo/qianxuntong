from flask import (
    Flask,
    request,
    render_template,
    redirect,
    url_for,
    jsonify,
)
import os
from werkzeug.utils import secure_filename
from sqlalchemy import and_
from utils.signer import insert_signatures_into_pdf
import uuid
import json
import base64
import re
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import Text
from datetime import datetime
from pytz import timezone
from urllib.parse import quote
from flask_login import UserMixin
from flask_login import (
    LoginManager,
    login_user,
    logout_user,
    login_required,
    current_user,
)
from werkzeug.security import generate_password_hash, check_password_hash
import tempfile
import platform
from flask import request, send_file, jsonify
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from flask_migrate import Migrate
import shutil
from PIL import Image
from reportlab.pdfgen import canvas as pdf_canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from docx2pdf import convert as docx2pdf_convert
import fitz  # PyMuPDF
import pdfplumber
from pdf2image import convert_from_path
import pytesseract
from openai import OpenAI
import hashlib
import subprocess

# ========== Flask app与数据库配置 ==========
app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 30 * 1024 * 1024  # 限制为30MB
app.config["SECRET_KEY"] = "very-secret-key-123456"
app.config["UPLOAD_FOLDER"] = "static/uploads"
app.config["FINAL_FOLDER"] = "static/final"
app.config["SIGN_URL"] = "http://127.0.0.1:5000/sign/"

# 读取数据库连接（支持环境变量）
app.config["SQLALCHEMY_DATABASE_URI"] = os.getenv(
    "DATABASE_URL", "mysql+pymysql://root:qxt123456@db:3306/qianxuntong?charset=utf8mb4"
)
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
db = SQLAlchemy(app)
migrate = Migrate(app, db)
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
os.makedirs(app.config["FINAL_FOLDER"], exist_ok=True)

CHINA_TZ = timezone("Asia/Shanghai")  # 定义中国时区


# ========== 数据库模型定义 ==========
# --- 用户表 ---
class User(UserMixin, db.Model):
    __table_args__ = {"mysql_charset": "utf8mb4", "mysql_collate": "utf8mb4_unicode_ci"}
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(255), nullable=False)


# --- 员工表 ---
class Employee(db.Model):
    __table_args__ = {"mysql_charset": "utf8mb4", "mysql_collate": "utf8mb4_unicode_ci"}
    id = db.Column(db.Integer, primary_key=True)  # 物理主键
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    name = db.Column(db.String(50), nullable=False)


# --- 签名任务表 ---
class SignatureTask(db.Model):
    __table_args__ = {"mysql_charset": "utf8mb4", "mysql_collate": "utf8mb4_unicode_ci"}
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    task_id = db.Column(db.String(64), unique=True, nullable=False)
    title = db.Column(db.String(255), nullable=False)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(CHINA_TZ))
    is_completed = db.Column(db.Boolean, default=False)
    employee_ids = db.Column(Text, default="[]")  # ✅ 新增字段（存为 JSON 字符串）
    quiz_required = db.Column(db.Boolean, default=False)

    def get_employee_ids(self):
        try:
            return json.loads(self.employee_ids)
        except:
            return []


# --- 员工签名状态表 ---
class SignatureStatus(db.Model):
    __table_args__ = {"mysql_charset": "utf8mb4", "mysql_collate": "utf8mb4_unicode_ci"}
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    task_id = db.Column(
        db.String(64), db.ForeignKey("signature_task.task_id"), nullable=False
    )
    employee_id = db.Column(db.Integer, db.ForeignKey("employee.id"), nullable=False)
    signed = db.Column(db.Boolean, default=False)
    signed_at = db.Column(db.DateTime)
    quiz_passed = db.Column(db.Boolean, default=False)


# --- 签名任务题库（答题）表 ---
class QuizQuestion(db.Model):
    __table_args__ = {"mysql_charset": "utf8mb4", "mysql_collate": "utf8mb4_unicode_ci"}
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    task_id = db.Column(
        db.String(64), db.ForeignKey("signature_task.task_id"), nullable=False
    )
    content = db.Column(db.Text, nullable=False)  # 题干内容
    options = db.Column(
        db.Text, nullable=False
    )  # JSON 字符串，例如：["选项A", "选项B", "选项C"]
    correct_answers = db.Column(
        db.Text, nullable=False
    )  # JSON 字符串，例如：[0] 或 [0, 2] 表示正确答案下标
    multiple = db.Column(db.Boolean, default=False)  # 是否为多选题


# --- 签名框区域表 ---
class SignatureBox(db.Model):
    __table_args__ = {"mysql_charset": "utf8mb4", "mysql_collate": "utf8mb4_unicode_ci"}
    id = db.Column(db.Integer, primary_key=True)
    task_id = db.Column(
        db.String(64), db.ForeignKey("signature_task.task_id"), nullable=False
    )
    employee_id = db.Column(db.Integer, db.ForeignKey("employee.id"), nullable=False)
    page = db.Column(db.Integer, nullable=False)
    left = db.Column(db.Float, nullable=False)
    top = db.Column(db.Float, nullable=False)
    width = db.Column(db.Float, nullable=False)
    height = db.Column(db.Float, nullable=False)
    signed = db.Column(db.Boolean, default=False)
    image = db.Column(db.Text)  # base64字符串，签名前可为空
    preview_width = db.Column(db.Float)
    preview_height = db.Column(db.Float)


# --- 培训材料表 ---
class TrainingMaterial(db.Model):
    __tablename__ = "training_material"
    __table_args__ = {"mysql_charset": "utf8mb4", "mysql_collate": "utf8mb4_unicode_ci"}
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    title = db.Column(db.String(255), nullable=False)
    description = db.Column(db.Text)
    file_path = db.Column(db.String(255), nullable=False)
    md5 = db.Column(db.String(32), unique=True, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    text_content = db.Column(db.Text)


# --- 培训题库表 ---
class TrainingQuestion(db.Model):
    __tablename__ = "training_question"
    __table_args__ = {"mysql_charset": "utf8mb4", "mysql_collate": "utf8mb4_unicode_ci"}
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    material_id = db.Column(
        db.Integer, db.ForeignKey("training_material.id"), nullable=False
    )
    content = db.Column(db.Text, nullable=False)
    options = db.Column(db.Text, nullable=False)  # JSON: ["A", "B", ...]
    correct_answers = db.Column(db.Text, nullable=False)  # JSON: [0,2]
    multiple = db.Column(db.Boolean, default=False)


# --- 培训任务表 ---
class TrainingTask(db.Model):
    __tablename__ = "training_task"
    __table_args__ = {"mysql_charset": "utf8mb4", "mysql_collate": "utf8mb4_unicode_ci"}
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    title = db.Column(db.String(200))
    material_id = db.Column(db.Integer, db.ForeignKey("training_material.id"))
    description = db.Column(db.Text)
    deadline = db.Column(db.Date)
    created_at = db.Column(db.DateTime, default=datetime.now)
    max_attempts = db.Column(db.Integer, default=1)  # 最大答题次数
    pass_score_ratio = db.Column(db.Float, default=0.8)  # 通过所需正确率（如0.8=80%）


# --- 培训任务-员工关联表 ---
class TrainingTaskEmployee(db.Model):
    __tablename__ = "training_task_employee"
    __table_args__ = {"mysql_charset": "utf8mb4", "mysql_collate": "utf8mb4_unicode_ci"}
    id = db.Column(db.Integer, primary_key=True)
    task_id = db.Column(db.Integer, db.ForeignKey("training_task.id"))
    employee_id = db.Column(db.Integer, db.ForeignKey("employee.id"))
    status = db.Column(db.String(10), default="未完成")
    score = db.Column(db.Float)
    submit_time = db.Column(db.DateTime)
    attempts = db.Column(db.Integer, default=0)  # 已答题次数
    is_passed = db.Column(db.Boolean, default=False)  # 是否已通过


# --- 培训答题历史表 ---
class TrainingAnswerHistory(db.Model):
    __tablename__ = "training_answer_history"
    __table_args__ = {"mysql_charset": "utf8mb4", "mysql_collate": "utf8mb4_unicode_ci"}
    id = db.Column(db.Integer, primary_key=True)
    task_id = db.Column(db.Integer, db.ForeignKey("training_task.id"))
    employee_id = db.Column(db.Integer, db.ForeignKey("employee.id"))
    attempt_num = db.Column(db.Integer)  # 第几次答题
    score = db.Column(db.Float)
    is_passed = db.Column(db.Boolean)
    submit_time = db.Column(db.DateTime, default=datetime.now)


# --- 培训任务制表（记录表） ---
class TrainingRecord(db.Model):
    __tablename__ = "training_record"
    __table_args__ = {"mysql_charset": "utf8mb4", "mysql_collate": "utf8mb4_unicode_ci"}
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    task_id = db.Column(
        db.Integer, db.ForeignKey("training_task.id"), unique=True, nullable=False
    )  # 每个任务一条
    station = db.Column(db.String(64))
    title = db.Column(db.String(128))
    time = db.Column(db.String(64))
    place = db.Column(db.String(64))
    trainer = db.Column(db.String(64))
    employees = db.Column(db.Text)
    summary = db.Column(db.Text)
    result = db.Column(db.Text)
    note = db.Column(db.Text)
    date = db.Column(db.String(32))
    review_date = db.Column(db.String(32))
    updated_at = db.Column(
        db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow
    )


# ===============================
# 1. 登录/注册/用户体系（系统通用区）
# ===============================
login_manager = LoginManager(app)
login_manager.login_view = "login"


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]
        if User.query.filter_by(username=username).first():
            return render_template("register.html", error="用户名已存在")
        user = User(username=username, password_hash=generate_password_hash(password))
        db.session.add(user)
        db.session.commit()
        # 注册成功，重定向到登录页面，并带参数
        return redirect(url_for("login", msg="注册成功，请登录"))
    return render_template("register.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    # 优先从GET参数获取提示信息
    msg = request.args.get("msg")
    error = None
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]
        user = User.query.filter_by(username=username).first()
        if user and check_password_hash(user.password_hash, password):
            login_user(user)
            return redirect(url_for("index"))
        error = "用户名或密码错误"
    return render_template("login.html", error=error, msg=msg)


@app.route("/logout")
@login_required
def logout():
    logout_user()
    return redirect(url_for("login"))


# ===============================
# 2. 员工管理区（全系统基础数据）
# ===============================
# 新增员工API，前端表单调用
@app.route("/employee/new", methods=["POST"])
@login_required
def add_employee():
    name = request.form["name"]
    new_emp = Employee(name=name, user_id=current_user.id)
    db.session.add(new_emp)
    db.session.commit()
    return jsonify({"status": "success", "id": new_emp.id, "name": new_emp.name})


# 删除员工API，同时级联删除相关记录
@app.route("/employee/delete/<int:id>", methods=["POST"])
@login_required
def delete_employee(id):
    emp = Employee.query.filter_by(user_id=current_user.id, id=id).first()
    if emp:
        # 先删所有有关的子表数据
        SignatureStatus.query.filter_by(employee_id=emp.id).delete()
        SignatureBox.query.filter_by(employee_id=emp.id).delete()
        TrainingTaskEmployee.query.filter_by(employee_id=emp.id).delete()
        TrainingAnswerHistory.query.filter_by(employee_id=emp.id).delete()

        db.session.delete(emp)
        db.session.commit()
        return jsonify({"status": "success"})
    return jsonify({"status": "not_found"}), 404


# ===============================
# 3. 签名任务系统（上传、配置、签名、进度管理）
# ===============================
# 首页与签名任务主入口
@app.route("/", methods=["GET", "POST"])
@login_required
def index():
    ALLOWED_EXTENSIONS = {"pdf", "doc", "docx", "jpg", "jpeg", "png"}

    def allowed_file(filename):
        return (
            "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS
        )

    if request.method == "POST":
        title = request.form.get("title", "").strip()
        file = request.files.get("file")
        employee_ids = [int(eid) for eid in request.form.getlist("employee_ids")]
        quiz_required = bool(request.form.get("quiz_required"))

        if not title or not file or not allowed_file(file.filename):
            return "请上传 PDF/Word/图片文件并填写标题", 400

        if not employee_ids:
            return "请至少选择一个员工", 400

        ext = file.filename.rsplit(".", 1)[1].lower()
        filename = secure_filename(file.filename)
        task_id = str(uuid.uuid4())
        user_folder = os.path.join(app.config["UPLOAD_FOLDER"], str(current_user.id))
        os.makedirs(user_folder, exist_ok=True)

        # 统一 PDF 存储路径
        pdf_filename = f"{task_id}.pdf"
        pdf_save_path = os.path.join(user_folder, pdf_filename)

        # ===== 文件类型判断与转换 =====
        if ext == "pdf":
            file.save(pdf_save_path)
        elif ext in {"doc", "docx"}:
            # 保存 word 文件再转换
            word_path = os.path.join(user_folder, f"{task_id}.{ext}")
            file.save(word_path)
            sysplat = platform.system()
            try:
                if sysplat == "Linux":
                    os.system(
                        f'libreoffice --headless --convert-to pdf "{word_path}" --outdir "{user_folder}"'
                    )
                    real_pdf_path = os.path.join(
                        user_folder, filename.rsplit(".", 1)[0] + ".pdf"
                    )
                    # 尝试找转换后的 PDF，如果没找到直接用 task_id.pdf
                    if os.path.exists(real_pdf_path):
                        shutil.move(real_pdf_path, pdf_save_path)
                    elif os.path.exists(word_path.replace(f".{ext}", ".pdf")):
                        shutil.move(word_path.replace(f".{ext}", ".pdf"), pdf_save_path)
                    else:
                        return "Word转PDF失败（未生成PDF）", 500
                else:
                    docx2pdf_convert(word_path, pdf_save_path)
            except Exception as e:
                return f"Word转PDF失败: {e}", 500
            os.remove(word_path)
        elif ext in {"jpg", "jpeg", "png"}:
            # 单图片转 PDF
            img = Image.open(file.stream)
            img = img.convert("RGB")
            img.save(pdf_save_path, "PDF")
        else:
            return "不支持的文件类型", 400

        # ====== 签名任务保存，PDF 参与后续流程 ======
        task = SignatureTask(
            user_id=current_user.id,
            task_id=task_id,
            title=title,
            employee_ids=json.dumps(employee_ids),
            quiz_required=quiz_required,
        )
        db.session.add(task)
        db.session.commit()

        # 保存员工签名状态
        for emp_id in employee_ids:
            emp = Employee.query.filter_by(id=emp_id, user_id=current_user.id).first()
            if emp:
                status = SignatureStatus(
                    user_id=current_user.id,
                    task_id=task_id,
                    employee_id=emp.id,
                    signed=False,
                )
                db.session.add(status)
        db.session.commit()

        # =========== 题库校验 =============
        question_keys = [
            k
            for k in request.form.keys()
            if re.match(r"questions\[(\d+)]\[content]", k)
        ]
        question_indexes = sorted(
            {
                int(re.findall(r"questions\[(\d+)]\[content]", k)[0])
                for k in question_keys
            }
        )

        for i in question_indexes:
            content = request.form.get(f"questions[{i}][content]", "").strip()
            if not content:
                continue
            options = request.form.getlist(f"questions[{i}][options][]")
            answer = request.form.get(f"questions[{i}][answers]")
            if answer is None or str(answer).strip() == "":
                return f"第{i+1}题未选择正确答案，请返回完善后再提交", 400
            if len(options) < 2:
                return f"第{i+1}题选项不足2个", 400
            try:
                correct_answers = [int(answer)]
            except Exception:
                return f"第{i+1}题正确答案格式不正确", 400

            quiz = QuizQuestion(
                user_id=current_user.id,
                task_id=task_id,
                content=content,
                options=json.dumps(options, ensure_ascii=False),
                correct_answers=json.dumps(correct_answers),
            )
            db.session.add(quiz)
        db.session.commit()

        return redirect(url_for("preview", task_id=task_id))

    # ✅ 获取签名任务及进度信息
    tasks = (
        SignatureTask.query.filter_by(user_id=current_user.id)
        .order_by(SignatureTask.created_at.desc())
        .all()
    )
    task_info = []

    for task in tasks:
        statuses = SignatureStatus.query.filter_by(
            user_id=current_user.id, task_id=task.task_id
        ).all()
        signed_count = sum(1 for s in statuses if s.signed)
        total_count = len(statuses)
        task_info.append(
            {
                "id": task.id,
                "task_id": task.task_id,
                "title": task.title,
                "created_at": task.created_at.strftime("%Y-%m-%d %H:%M:%S"),
                "progress": f"{signed_count}/{total_count}",
                "is_completed": task.is_completed,
            }
        )

    employees = (
        Employee.query.filter_by(user_id=current_user.id)
        .order_by(Employee.id.desc())
        .all()
    )
    materials = (
        TrainingMaterial.query.filter_by(user_id=current_user.id)
        .order_by(TrainingMaterial.created_at.desc())
        .all()
    )
    return render_template(
        "index.html", tasks=task_info, employees=employees, materials=materials
    )


# 预览签名任务主页面，主要传递PDF/签名区域/员工数据
@app.route("/preview/<task_id>")
@login_required
def preview(task_id):

    boxes = SignatureBox.query.filter_by(task_id=task_id).all()
    sign_ready = len(boxes) > 0
    sign_link = app.config["SIGN_URL"] + task_id if sign_ready else None

    employees_raw = Employee.query.filter_by(user_id=current_user.id).all()
    employees = [
        {
            "id": emp.id,
            "name": emp.name,
        }
        for emp in employees_raw
    ]

    task = SignatureTask.query.filter_by(
        user_id=current_user.id, task_id=task_id
    ).first()

    user_folder = os.path.join(app.config["UPLOAD_FOLDER"], str(current_user.id))
    uploaded_filename = next(
        (f for f in os.listdir(user_folder) if f.startswith(task_id)),
        None,
    )

    if not uploaded_filename:
        return (
            render_template(
                "message.html",
                title="未找到文件",
                msg="未找到上传的 PDF 文件",
                btn_text="关闭页面",
                back_url=url_for("index"),
            ),
            404,
        )

    encoded_filename = quote(uploaded_filename)  # ✅ URL 编码

    return render_template(
        "preview.html",
        task_id=task_id,
        sign_link=sign_link,
        sign_ready=sign_ready,
        employees=employees,
        record=task,  # ✅ 加入传递
        encoded_title=encoded_filename,
    )


# 保存签名区域，前端配置签名框后调用
@app.route("/save_box/<task_id>", methods=["POST"])
@login_required
def save_box(task_id):
    box_data = request.get_json()  # 这是前端传来的所有box列表

    # 先删除本任务下的所有老的box（防止反复保存产生重复/脏数据）
    SignatureBox.query.filter_by(task_id=task_id).delete()
    db.session.commit()

    # 批量插入每个box
    for box in box_data:
        try:
            # ⭐⭐ 这里做一次 id -> 主键 id 的映射 ⭐⭐
            emp = Employee.query.filter_by(
                user_id=current_user.id, id=int(box["employee_id"])
            ).first()
            if not emp:
                continue  # 跳过找不到的

            new_box = SignatureBox(
                task_id=task_id,
                employee_id=emp.id,  # 用主键 id!!
                page=int(box["page"]),
                left=float(box["left"]),
                top=float(box["top"]),
                width=float(box["width"]),
                height=float(box["height"]),
                signed=bool(box.get("signed", False)),
                image=box.get("image"),  # 刚配置时一般为 None
                preview_width=float(box.get("preview_width", 0)),
                preview_height=float(box.get("preview_height", 0)),
            )
            db.session.add(new_box)
        except Exception as e:
            print("❌ 保存 box 时出错", box, e)
    db.session.commit()

    return jsonify({"status": "success"})


# 签名主页面（员工签名入口，批量签）
@app.route("/sign/<task_id>")
def sign_page(task_id):
    boxes = SignatureBox.query.filter_by(task_id=task_id).all()
    if not boxes:
        return "签名区域未配置"

    # 如果模板用 dict，可以转一下；如果模板直接用对象列表也可以
    box_list = [
        {
            "id": box.id,
            "employee_id": box.employee_id,
            "page": int(box.page),
            "left": float(box.left),
            "top": float(box.top),
            "width": float(box.width),
            "height": float(box.height),
            "signed": bool(box.signed),
            "image": box.image,
            "preview_width": float(box.preview_width) if box.preview_width else None,
            "preview_height": float(box.preview_height) if box.preview_height else None,
        }
        for box in boxes
    ]

    return render_template("sign.html", task_id=task_id, boxes=box_list)


# 签名提交API，保存签名图、状态并自动合成PDF
@app.route("/submit_sign/<task_id>", methods=["POST"])
def submit_sign(task_id):
    task = SignatureTask.query.filter_by(
        user_id=current_user.id, task_id=task_id
    ).first()
    if task and task.is_completed:
        return (
            render_template(
                "message.html",
                title="无法签名",
                msg="该签名任务已完成，无法继续签名",
                btn_text="返回首页",
            ),
            403,
        )

    data = request.get_json()

    if not data or not isinstance(data, list):
        return jsonify({"status": "error", "msg": "缺少签名数据"})

    employee_id = data[0].get("employee_id")
    if not employee_id:
        return jsonify({"status": "error", "msg": "缺少签名人信息"})

    updated = False
    for item in data:
        box = SignatureBox.query.filter(
            SignatureBox.task_id == task_id,
            SignatureBox.employee_id == int(item["employee_id"]),
            SignatureBox.page == int(item["page"]),
            db.func.abs(SignatureBox.left - float(item["left"])) < 0.01,
            db.func.abs(SignatureBox.top - float(item["top"])) < 0.01,
        ).first()
        if box and not box.signed:
            box.signed = True
            box.image = item.get("image")
            box.preview_width = item.get("preview_width")
            box.preview_height = item.get("preview_height")
            updated = True

    if not updated:
        return jsonify({"status": "error", "msg": "未匹配到对应签名区域，或已签名"})

    db.session.commit()

    # 更新数据库中的签名状态
    sig_status = SignatureStatus.query.filter_by(
        user_id=current_user.id, task_id=task_id, employee_id=int(employee_id)
    ).first()
    if sig_status and not sig_status.signed:
        sig_status.signed = True
        db.session.commit()

    # 检查是否所有人都签完，才触发合成 PDF
    all_statuses = SignatureStatus.query.filter_by(
        user_id=current_user.id, task_id=task_id
    ).all()
    if all(s.signed for s in all_statuses):

        user_folder = os.path.join(app.config["UPLOAD_FOLDER"], str(current_user.id))
        pdf_path = next(
            (f for f in os.listdir(user_folder) if f.startswith(task_id)),
            None,
        )
        if not pdf_path:
            return jsonify({"status": "error", "msg": "PDF 未找到"})

        full_pdf_path = os.path.join(user_folder, pdf_path)
        final_user_folder = os.path.join(
            app.config["FINAL_FOLDER"], str(current_user.id)
        )
        os.makedirs(final_user_folder, exist_ok=True)
        output_pdf_path = os.path.join(final_user_folder, f"{task_id}_signed.pdf")

        # 构建签名图列表（已签名的）
        signature_boxes = SignatureBox.query.filter_by(
            task_id=task_id, signed=True
        ).all()
        signature_images = []
        for box in signature_boxes:
            if box.image:
                image_base64 = (
                    box.image.split(",")[1] if "," in box.image else box.image
                )
                signature_images.append(
                    {
                        "page": int(box.page),
                        "left": float(box.left),
                        "top": float(box.top),
                        "width": float(box.width),
                        "height": float(box.height),
                        "image_bytes": base64.b64decode(image_base64),
                        "preview_width": (
                            float(box.preview_width) if box.preview_width else 1240
                        ),
                        "preview_height": (
                            float(box.preview_height) if box.preview_height else 1754
                        ),
                    }
                )

        insert_signatures_into_pdf(full_pdf_path, output_pdf_path, signature_images)

        # 标记任务完成
        task.is_completed = True
        db.session.commit()

    return jsonify(
        {"status": "success", "redirect": url_for("sign_submitted", task_id=task_id)}
    )


@app.route("/sign_submitted/<task_id>")
def sign_submitted(task_id):
    return render_template("sign_submitted.html", task_id=task_id)


# 签名邀请、分员工签名入口
@app.route("/invite/<task_id>")
@login_required
def invite_page(task_id):
    # 查找该任务下所有签名区域（SignatureBox）
    boxes = SignatureBox.query.filter_by(task_id=task_id).all()
    if not boxes:
        return "签名区域未配置"

    # 提取所有涉及到的员工ID
    employee_ids = list({box.employee_id for box in boxes})

    # 获取这些员工的信息
    employees = (
        Employee.query.filter_by(user_id=current_user.id)
        .filter(Employee.id.in_(employee_ids))
        .all()
    )

    # 已签名的员工ID集合
    signed_ids = {box.employee_id for box in boxes if box.signed}

    base_url = request.url_root.rstrip("/")

    return render_template(
        "sign_invite.html",
        task_id=task_id,
        sign_url=f"{base_url}/sign_select/{task_id}",
        employees=employees,
        signed_ids=signed_ids,
    )


# Canvas签名绘制页面
@app.route("/sign_canvas/<task_id>", defaults={"employee_id": None})
@app.route("/sign_canvas/<task_id>/<int:employee_id>")
def sign_canvas(task_id, employee_id):
    return render_template("sign_canvas.html", task_id=task_id, employee_id=employee_id)


# 删除签名任务，级联删除所有相关文件和记录
@app.route("/delete_record/<task_id>", methods=["POST"])
@login_required
def delete_record(task_id):
    task = SignatureTask.query.filter_by(
        user_id=current_user.id, task_id=task_id
    ).first()
    if task:
        # 删除关联签名状态
        SignatureStatus.query.filter_by(
            user_id=current_user.id, task_id=task_id
        ).delete()

        # 删除所有签名区域
        SignatureBox.query.filter_by(task_id=task_id).delete()

        # 删除题库
        QuizQuestion.query.filter_by(user_id=current_user.id, task_id=task_id).delete()

        # 删除任务本身
        db.session.delete(task)
        db.session.commit()

        # 删除上传的 PDF 文件
        user_folder = os.path.join(app.config["UPLOAD_FOLDER"], str(current_user.id))
        if os.path.exists(user_folder):
            for fname in os.listdir(user_folder):
                if fname.startswith(task_id):
                    os.remove(os.path.join(user_folder, fname))

        # 删除合成后的 PDF 文件
        final_user_folder = os.path.join(
            app.config["FINAL_FOLDER"], str(current_user.id)
        )
        final_path = os.path.join(final_user_folder, f"{task_id}_signed.pdf")
        if os.path.exists(final_path):
            os.remove(final_path)

        return jsonify({"status": "success"})

    return jsonify({"status": "not_found"}), 404


# 签名前选择员工身份页面
@app.route("/sign_select/<task_id>", methods=["GET", "POST"])
def sign_select(task_id):
    task = SignatureTask.query.filter_by(
        user_id=current_user.id, task_id=task_id
    ).first()
    if task and task.is_completed:
        return (
            render_template(
                "message.html",
                title="无法签名",
                msg="该签名任务已完成，无法继续签名",
                btn_text="关闭页面",
            ),
            403,
        )

    if request.method == "POST":
        employee_id = request.form.get("employee_id")
        return redirect(
            url_for("sign_page_employee", task_id=task_id, employee_id=employee_id)
        )

    # 获取当前任务所有涉及到的员工ID（查 SignatureBox）
    box_emps = SignatureBox.query.filter_by(task_id=task_id).all()
    employee_ids = set(b.employee_id for b in box_emps)
    employees = (
        Employee.query.filter_by(user_id=current_user.id)
        .filter(Employee.id.in_(employee_ids))
        .all()
    )

    return render_template("sign_select.html", task_id=task_id, employees=employees)


# 针对指定员工的签名页
@app.route("/sign/<task_id>/<int:employee_id>")
def sign_page_employee(task_id, employee_id):
    # ✅ 获取任务
    task = SignatureTask.query.filter_by(
        user_id=current_user.id, task_id=task_id
    ).first()
    if not task:
        return (
            render_template(
                "message.html",
                title="任务不存在",
                msg="签名任务不存在",
                btn_text="关闭页面",
            ),
            404,
        )

    # ✅ 检查是否任务已完成
    status = SignatureStatus.query.filter_by(
        user_id=current_user.id, task_id=task_id, employee_id=employee_id
    ).first()
    if status and status.signed:
        return (
            render_template(
                "message.html",
                title="已签名",
                msg="您已完成签名，无法再次签名",
                btn_text="关闭页面",
            ),
            403,
        )

    # ✅ 获取 quiz 状态，前端判断是否允许签名，不再强制跳 quiz
    quiz_passed = status.quiz_passed if status else False

    # ✅ 直接查 SignatureBox 表，获取当前员工所有签名区域
    boxes_raw = SignatureBox.query.filter_by(
        task_id=task_id, employee_id=employee_id
    ).all()
    filtered_boxes = []
    for box in boxes_raw:
        filtered_boxes.append(
            {
                "page": int(box.page),
                "left": float(box.left),
                "top": float(box.top),
                "width": float(box.width),
                "height": float(box.height),
                "signed": box.signed,
                "image": box.image,
                "preview_width": (
                    float(box.preview_width) if box.preview_width else None
                ),
                "preview_height": (
                    float(box.preview_height) if box.preview_height else None
                ),
                "employee_id": box.employee_id,
            }
        )

    # ✅ 获取上传的 PDF 文件名并 URL 编码
    user_folder = os.path.join(app.config["UPLOAD_FOLDER"], str(current_user.id))
    uploaded_filename = next(
        (f for f in os.listdir(user_folder) if f.startswith(task_id)),
        None,
    )
    if not uploaded_filename:
        return (
            render_template(
                "message.html",
                title="未找到文件",
                msg="未找到上传的 PDF 文件",
                btn_text="关闭页面",
            ),
            404,
        )

    encoded_title = quote(uploaded_filename)
    employee = Employee.query.filter_by(user_id=current_user.id, id=employee_id).first()
    employee_name = employee.name if employee else ""

    return render_template(
        "sign.html",
        task_id=task_id,
        boxes=filtered_boxes,
        employee_id=employee_id,
        employee_name=employee_name,
        encoded_title=encoded_title,
        quiz_required=task.quiz_required,
        quiz_passed=quiz_passed,
    )


def extract_text_from_pdf(pdf_path):
    """
    从PDF文件中提取文本，优先使用电子文本方式，最后使用OCR兜底。
    支持中文和英文PDF自动识别。
    """
    # 第一优先：用 PyMuPDF 直接读取PDF文本（适合电子PDF）
    print(f"[DEBUG] 开始解析PDF：{pdf_path}")
    try:
        doc = fitz.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        print(f"[DEBUG] PyMuPDF 提取文本长度: {len(text)}")
        # 提取到较多文本则直接返回
        if len(text.strip()) > 30:
            print("[DEBUG] PyMuPDF 提取成功")
            return text
    except Exception:
        print(f"[DEBUG] PyMuPDF 解析异常: {e}")
        pass

    # 第二优先：用 pdfplumber 尝试读取文本
    try:
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        print(f"[DEBUG] pdfplumber 提取文本长度: {len(text)}")
        if len(text.strip()) > 30:
            print("[DEBUG] pdfplumber 提取成功")
            return text
    except Exception:
        print(f"[DEBUG] pdfplumber 解析异常: {e}")
        pass

    # 最后兜底：PDF转图片后用OCR识别，适合扫描件/图片型PDF
    try:
        images = convert_from_path(pdf_path)
        text = ""
        for img in images:
            text += pytesseract.image_to_string(img, lang="chi_sim+eng") + "\n"
        print(f"[DEBUG] OCR 提取文本长度: {len(text)}")
        return text
    except Exception:
        print(f"[DEBUG] OCR 解析异常: {e}")
        return ""


@app.route("/api/pdf2text", methods=["POST"])
def pdf2text():
    """
    接收前端上传的PDF文件，提取出可识别的文本内容（最多3万字），返回json。
    """
    if "file" not in request.files:
        return jsonify({"status": "fail", "msg": "未上传文件"}), 400
    file = request.files["file"]
    if not file.filename.lower().endswith(".pdf"):
        return jsonify({"status": "fail", "msg": "请上传PDF文件"}), 400

    # 文件临时保存路径
    save_path = os.path.join("/tmp", file.filename)
    file.save(save_path)

    # 提取文本
    text = extract_text_from_pdf(save_path)
    os.remove(save_path)

    # 未识别出有效文本
    if not text.strip():
        return jsonify({"status": "fail", "msg": "未能识别出文本"}), 200

    # 成功返回，文本过长则截断3万字
    return jsonify({"status": "success", "text": text[:30000]})


def extract_jsons(s):
    # 优先找[...]
    arr_match = re.search(r"(\[[\s\S]+\])", s)
    if arr_match:
        try:
            return json.loads(arr_match.group(1))
        except Exception:
            pass
    # fallback：找所有 {...}
    obj_matches = re.findall(r"\{[\s\S]*?\}", s)
    if obj_matches:
        res = []
        for obj in obj_matches:
            try:
                res.append(json.loads(obj))
            except Exception:
                continue
        if res:
            return res
    # fallback: 全部失败
    raise ValueError("AI未输出合法JSON结构")


DEEPSEEK_KEY = "sk-6a104b306d4d41378825c18da83b0e6b"
client = OpenAI(api_key=DEEPSEEK_KEY, base_url="https://api.deepseek.com")


@app.route("/api/sign_ai_generate_questions", methods=["POST"])
@login_required
def sign_ai_generate_questions():
    """
    专供签名系统AI生成题目。
    前端传：text, type, count, level
    type: single(单选题), judge(判断题)  (默认single)
    """
    data = request.get_json()
    text = data.get("text", "")
    q_type = data.get("type", "single")  # 'single' or 'judge'
    count = int(data.get("count", 3))
    level = data.get("level", "easy")  # easy/deep

    if not text or len(text.strip()) < 40:
        return jsonify({"status": "fail", "msg": "请先上传PDF并确保内容足够"}), 200

    # 拼接prompt
    if q_type == "single":
        prompt = f"""
        请根据以下材料内容，自动生成{count}道“单选题”（每题4个选项，且只有1个正确答案），题目难度为“{level}”。
        严格输出如下格式的JSON数组（无说明）：
        [
        {{
            "content": "题干内容",
            "options": ["A选项", "B选项", "C选项", "D选项"],
            "answer": 0
        }}
        ]
        材料内容如下：
        -----------------
        {text}
        -----------------
        """
    elif q_type == "judge":
        prompt = f"""
        请根据以下材料内容，自动生成{count}道“判断题”。每题答案为“正确”或“错误”二选一，题目难度为“{level}”。
        严格输出如下格式的JSON数组（无说明）：
        [
        {{
            "content": "题干内容",
            "answer": 0
        }}
        ]
        材料内容如下：
        -----------------
        {text}
        -----------------
        """
    else:
        return jsonify({"status": "fail", "msg": "暂不支持的题型type"}), 200

    # AI调用
    try:
        response = client.chat.completions.create(
            model="deepseek-reasoner",
            messages=[
                {
                    "role": "system",
                    "content": "你是一个擅长出题的老师，只输出 JSON 数组本身，不要有任何注释、说明或多余内容。",
                },
                {"role": "user", "content": prompt},
            ],
            response_format={"type": "json_object"},
            max_tokens=32000,
            temperature=1.0,
            stream=False,
        )
        output = response.choices[0].message.content.strip()
        questions = extract_jsons(output)

        # 数组形式容错
        if isinstance(questions, dict) and "questions" in questions:
            questions = questions["questions"]

        # 补充字段适配
        if q_type == "single":
            for q in questions:
                if "correct" in q:
                    q["answer"] = q["correct"]
                if "options" not in q:
                    q["options"] = ["A", "B", "C", "D"]
        elif q_type == "judge":
            for q in questions:
                q["options"] = ["正确", "错误"]
                if "correct" in q:
                    q["answer"] = q["correct"]
                if "正确" in q:
                    q["answer"] = q["正确"]
                try:
                    q["answer"] = int(q["answer"])
                except Exception:
                    q["answer"] = 0

        return jsonify({"status": "success", "questions": questions})

    except Exception as e:
        import traceback

        traceback.print_exc()
        code = 500
        msg = f"AI生成题目失败: {str(e)}"

        if hasattr(e, "status_code"):
            code = e.status_code
        elif hasattr(e, "http_status"):
            code = e.http_status
        elif hasattr(e, "error") and hasattr(e.error, "code"):
            code = e.error.code
        elif hasattr(e, "response") and hasattr(e.response, "status_code"):
            code = e.response.status_code
        elif isinstance(e, int):
            code = e

        if "余额不足" in str(e):
            code = 402
            msg = "API账户余额不足，请充值后重试。"
        elif "认证失败" in str(e):
            code = 401
            msg = "API认证失败，请检查API Key设置。"

        return jsonify({"status": "fail", "msg": msg, "code": code}), 200


# 后端API：POST /api/ai_generate_questions
@app.route("/api/ai_generate_questions", methods=["POST"])
@login_required
def ai_generate_questions():
    """
    前端传: material_id, text, type, count, level(签名系统用)
    type: single(单选题), judge(判断题)
    """
    data = request.get_json()
    material_id = data.get("material_id")
    text = data.get("text", "")
    q_type = data.get("type", "single")  # 'single' or 'judge'
    count = int(data.get("count", 3))
    level = data.get("level", "easy")  # easy/deep

    # 获取材料的文本内容
    material = TrainingMaterial.query.filter_by(
        id=material_id, user_id=current_user.id
    ).first()
    if not material:
        return jsonify({"status": "fail", "msg": "材料不存在"}), 404

    # 获取材料的文本内容（直接从数据库中获取）
    text = material.text_content

    # 基本校验
    if not text or len(text.strip()) < 40:
        return jsonify({"status": "fail", "msg": "材料文本内容过少"}), 200

    # 新增：获取当前材料下所有已有题干，防止重复
    existed_questions = TrainingQuestion.query.filter_by(
        material_id=material_id, user_id=current_user.id
    ).all()
    existed_contents = [
        q.content.strip() for q in existed_questions if q.content.strip()
    ]
    existed_prompt = ""
    if existed_contents:
        existed_prompt = "请不要生成与以下题目内容重复的题目：\n" + "\n".join(
            [f"{i+1}. {c}" for i, c in enumerate(existed_contents)]
        )

    # 1. 拼接prompt（题型适配）
    if q_type == "single":
        prompt = f"""
        {existed_prompt}

        请根据以下材料内容，自动生成{count}道“单选题”（每题4个选项，且只有1个正确答案），题目难度为“{level}”。
        要求内容创新，不能与已有题目内容和问法重复。
        严格输出如下格式的JSON数组（无说明）：
        [
        {{
            "content": "题干内容",
            "options": ["A选项", "B选项", "C选项", "D选项"],
            "answer": 0
        }}
        ]
        材料内容如下：
        -----------------
        {text}
        -----------------
        """
    elif q_type == "judge":
        prompt = f"""
        {existed_prompt}

        请根据以下材料内容，自动生成{count}道“判断题”。每题答案为“正确”或“错误”二选一，题目难度为“{level}”。
        要求内容创新，不能与已有题目内容和问法重复。
        严格输出如下格式的JSON数组（无说明）：
        [
        {{
            "content": "题干内容",
            "answer": 0
        }}
        ]
        材料内容如下：
        -----------------
        {text}
        -----------------
        """
    else:
        return jsonify({"status": "fail", "msg": "暂不支持的题型type"}), 200

    # AI调用
    try:
        response = client.chat.completions.create(
            model="deepseek-reasoner",
            messages=[
                {
                    "role": "system",
                    "content": "你是一个擅长出题的老师，只输出 JSON 数组本身，不要有任何注释、说明或多余内容。",
                },
                {"role": "user", "content": prompt},
            ],
            response_format={"type": "json_object"},
            max_tokens=32000,
            temperature=1.0,
            stream=False,
        )
        output = response.choices[0].message.content.strip()
        questions = extract_jsons(output)

        # 数组形式容错
        if isinstance(questions, dict) and "questions" in questions:
            questions = questions["questions"]

        # 补充字段适配
        if q_type == "single":
            for q in questions:
                if "correct" in q:
                    q["answer"] = q["correct"]
                if "options" not in q:
                    q["options"] = ["A", "B", "C", "D"]
        elif q_type == "judge":
            for q in questions:
                q["options"] = ["正确", "错误"]
                if "correct" in q:
                    q["answer"] = q["correct"]
                if "正确" in q:
                    q["answer"] = q["正确"]
                try:
                    q["answer"] = int(q["answer"])
                except Exception:
                    q["answer"] = 0

        return jsonify({"status": "success", "questions": questions})

    except Exception as e:
        import traceback

        traceback.print_exc()
        code = 500
        msg = f"AI生成题目失败: {str(e)}"

        if hasattr(e, "status_code"):
            code = e.status_code
        elif hasattr(e, "http_status"):
            code = e.http_status
        elif hasattr(e, "error") and hasattr(e.error, "code"):
            code = e.error.code
        elif hasattr(e, "response") and hasattr(e.response, "status_code"):
            code = e.response.status_code
        elif isinstance(e, int):
            code = e

        if "余额不足" in str(e):
            code = 402
            msg = "API账户余额不足，请充值后重试。"
        elif "认证失败" in str(e):
            code = 401
            msg = "API认证失败，请检查API Key设置。"

        return jsonify({"status": "fail", "msg": msg, "code": code}), 200


# ===============================
# 4. 签名任务答题功能（签名前答题环节）
# ===============================
# 员工签名前答题入口，答题通过后允许签名
@app.route("/sign_quiz/<task_id>/<int:employee_id>", methods=["GET", "POST"])
def quiz_page(task_id, employee_id):
    task = SignatureTask.query.filter_by(
        user_id=current_user.id, task_id=task_id
    ).first()
    if not task:
        return (
            render_template(
                "message.html",
                title="任务不存在",
                msg="签名任务不存在",
                btn_text="关闭页面",
            ),
            404,
        )

    questions = QuizQuestion.query.filter_by(
        user_id=current_user.id, task_id=task_id
    ).all()

    if request.method == "POST":
        # 用 request.json 读取 AJAX 数据
        data = request.get_json()
        error_message = None
        for q in questions:
            correct = json.loads(q.correct_answers)
            input_name = f"q{q.id}"
            user_answer = data.get(input_name)
            try:
                user_answer = int(user_answer)
            except Exception:
                error_message = "答案格式错误"
                break
            if user_answer != correct[0]:
                error_message = "答题不通过，请重新作答"
                break

        if error_message:
            return jsonify({"success": False, "msg": error_message})

        # 更新当前员工 quiz_passed 状态
        status = SignatureStatus.query.filter_by(
            user_id=current_user.id, task_id=task_id, employee_id=employee_id
        ).first()
        if status:
            status.quiz_passed = True  # ✅ 标记通过
            db.session.commit()  # ✅ 写入数据库
        return jsonify(
            {
                "success": True,
                "redirect": url_for(
                    "sign_canvas", task_id=task_id, employee_id=employee_id
                ),
            }
        )

    # 👇 这一部分是修改的核心
    parsed_questions = []
    for q in questions:
        parsed_questions.append(
            {
                "id": q.id,
                "content": q.content,
                "options": json.loads(q.options),
                "correct_answers": json.loads(q.correct_answers),
            }
        )

    return render_template(
        "sign_quiz.html",
        task_id=task_id,
        employee_id=employee_id,
        questions=parsed_questions,
    )


# ===============================
# 5. 培训系统 —— 材料管理
# ===============================
# 上传材料API，存文件及记录（统一命名、规范化）
@app.route("/training_materials", methods=["POST"])
@login_required
def training_materials():
    user_folder = os.path.join("static/training_materials", str(current_user.id))
    os.makedirs(user_folder, exist_ok=True)
    title = request.form.get("title", "").strip()
    desc = request.form.get("description", "")
    file = request.files.get("file")

    ALLOWED_EXTS = {"pdf", "doc", "docx", "jpg", "jpeg", "png"}

    def allowed_file(filename):
        return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTS

    if not (title and file and allowed_file(file.filename)):
        return (
            jsonify({"status": "fail", "msg": "请上传PDF/Word/图片文件并填写标题"}),
            200,
        )

    # 计算文件的MD5值
    file_md5 = hashlib.md5(file.read()).hexdigest()
    file.seek(0)  # 重新设置文件指针，确保后续代码能正常读取文件

    # 检查文件是否已经存在
    existing_material = TrainingMaterial.query.filter_by(md5=file_md5).first()
    if existing_material:
        return (
            jsonify(
                {
                    "status": "exists",
                    "msg": "文件已存在",
                    "file_path": existing_material.file_path,
                    "text_content": existing_material.text_content,  # 如果已存在，返回文本内容
                }
            ),
            200,
        )

    ext = file.filename.rsplit(".", 1)[1].lower()

    # ========== 先插入材料表，获取唯一ID ==========
    mat = TrainingMaterial(
        user_id=current_user.id,
        title=title,
        description=desc,
        file_path="",  # 先占位，后面补
        md5=file_md5,  # 存储MD5值
    )
    db.session.add(mat)
    db.session.flush()  # 不commit可以拿到mat.id

    # ==== 规范文件名：{id}_{title}.pdf，过滤标题非法字符 ====
    safe_title = "".join(
        c for c in title if c.isalnum() or c in (" ", "_", "-", "（", "）")
    )
    if not safe_title:
        safe_title = str(mat.id)
    pdf_filename = f"{mat.id}_{safe_title}.pdf"
    pdf_save_path = os.path.join(user_folder, pdf_filename)

    try:
        if ext == "pdf":
            file.save(pdf_save_path)
        elif ext in {"doc", "docx"}:
            word_path = os.path.join(user_folder, f"{mat.id}_{safe_title}.{ext}")
            file.save(word_path)
            sysplat = platform.system()
            if sysplat == "Linux":
                os.system(
                    f'libreoffice --headless --convert-to pdf "{word_path}" --outdir "{user_folder}"'
                )
                # 有时转换文件名不是你想要的，这里兜底全部.pdf只取最新
                pdfs = sorted(
                    [f for f in os.listdir(user_folder) if f.endswith(".pdf")],
                    key=lambda fn: os.path.getmtime(os.path.join(user_folder, fn)),
                    reverse=True,
                )
                if pdfs:
                    shutil.move(os.path.join(user_folder, pdfs[0]), pdf_save_path)
                else:
                    return jsonify({"status": "fail", "msg": "Word转PDF失败"}), 500
                os.remove(word_path)
            else:
                docx2pdf_convert(word_path, pdf_save_path)
                os.remove(word_path)
        elif ext in {"jpg", "jpeg", "png"}:
            img = Image.open(file.stream)
            img = img.convert("RGB")
            img.save(pdf_save_path, "PDF")
        else:
            return jsonify({"status": "fail", "msg": "不支持的文件类型"}), 400
    except Exception as e:
        db.session.rollback()
        return jsonify({"status": "fail", "msg": f"文件转换出错：{e}"}), 500

    # 提取文本内容并存储到数据库
    text_content = extract_text_from_pdf(pdf_save_path)
    mat.text_content = text_content  # 保存提取的文本内容
    db.session.commit()  # 提交保存

    # ========== 更新file_path为规范路径并提交 ==========
    mat.file_path = os.path.join(str(current_user.id), pdf_filename)
    db.session.commit()

    return jsonify(
        {
            "status": "success",
            "mat": {
                "id": mat.id,
                "title": mat.title,
                "description": mat.description or "",
                "file_path": mat.file_path,
                "created_at": mat.created_at.strftime("%Y-%m-%d %H:%M"),
                "text_content": mat.text_content,  # 返回文本内容
            },
        }
    )


# 材料列表API，返回JSON
@app.route("/training_materials/list")
@login_required
def training_materials_list():
    mats = (
        TrainingMaterial.query.filter_by(user_id=current_user.id)
        .order_by(TrainingMaterial.created_at.desc())
        .all()
    )

    def mat2dict(m):
        return {
            "id": m.id,
            "title": m.title,
            "description": m.description or "",
            "created_at": m.created_at.strftime("%Y-%m-%d %H:%M"),
            "file_path": m.file_path,
            "text_content": m.text_content or "",  # Ensure text_content is included
        }

    return jsonify({"mats": [mat2dict(m) for m in mats]})


# 删除材料API
@app.route("/training_materials/delete/<int:mat_id>", methods=["POST"])
@login_required
def delete_material(mat_id):
    mat = TrainingMaterial.query.filter_by(id=mat_id, user_id=current_user.id).first()
    if mat:
        abs_path = os.path.join("static/training_materials", mat.file_path)
        if os.path.exists(abs_path):
            os.remove(abs_path)
        db.session.delete(mat)
        db.session.commit()
        return jsonify({"status": "success"})
    return jsonify({"status": "not_found"}), 404


@app.route("/training_materials/get_text/<int:material_id>")
@login_required
def get_material_text(material_id):
    mat = TrainingMaterial.query.filter_by(
        id=material_id, user_id=current_user.id
    ).first()
    if not mat:
        return jsonify({"status": "fail", "msg": "材料不存在"}), 404

    pdf_path = os.path.join("static/training_materials", mat.file_path)
    if not os.path.exists(pdf_path):
        return jsonify({"status": "fail", "msg": "文件不存在"}), 404

    try:
        text = extract_text_from_pdf(pdf_path)
        if not text or len(text.strip()) < 10:
            return jsonify({"status": "fail", "msg": "未能提取文本"}), 200
        return jsonify({"status": "success", "text": text[:30000]})
    except Exception as e:
        return jsonify({"status": "fail", "msg": f"提取异常: {e}"}), 500


# ===============================
# 6. 培训系统 —— 题库管理
# ===============================
# 加载某材料下题目
@app.route("/training_questions/list")
@login_required
def training_questions_list():
    material_id = request.args.get("material_id", type=int)
    questions = (
        TrainingQuestion.query.filter_by(
            user_id=current_user.id, material_id=material_id
        )
        .order_by(TrainingQuestion.id.desc())
        .all()
    )

    def q2dict(q):
        options = json.loads(q.options)
        # 判断题：options固定为["正确","错误"]
        if options == ["正确", "错误"]:
            qtype = "判断"
        elif q.multiple:
            qtype = "多选"
        else:
            qtype = "单选"
        return {
            "id": q.id,
            "content": q.content,
            "options": options,
            "correct_answers": json.loads(q.correct_answers),
            "multiple": q.multiple,
            "qtype": qtype,  # 新增类型字段
        }

    return jsonify({"questions": [q2dict(q) for q in questions]})


# 新建题目API
@app.route("/training_questions/new", methods=["POST"])
@login_required
def training_question_new():
    data = request.get_json()
    q = TrainingQuestion(
        user_id=current_user.id,
        material_id=data["material_id"],
        content=data["content"],
        options=json.dumps(data["options"]),
        correct_answers=json.dumps(data["correct_answers"]),
        multiple=data["multiple"],
    )
    db.session.add(q)
    db.session.commit()
    return jsonify({"status": "success"})


# 获取题目信息API
@app.route("/training_questions/get/<int:qid>")
@login_required
def training_question_get(qid):
    q = TrainingQuestion.query.filter_by(id=qid, user_id=current_user.id).first()
    if not q:
        return jsonify({"status": "not_found"})
    return jsonify(
        {
            "status": "success",
            "q": {
                "id": q.id,
                "content": q.content,
                "options": json.loads(q.options),
                "correct_answers": json.loads(q.correct_answers),
                "multiple": q.multiple,
            },
        }
    )


# 编辑题目API
@app.route("/training_questions/edit/<int:qid>", methods=["POST"])
@login_required
def training_question_edit(qid):
    q = TrainingQuestion.query.filter_by(id=qid, user_id=current_user.id).first()
    if not q:
        return jsonify({"status": "not_found"})
    data = request.get_json()
    q.content = data["content"]
    q.options = json.dumps(data["options"])
    q.correct_answers = json.dumps(data["correct_answers"])
    q.multiple = data["multiple"]
    db.session.commit()
    return jsonify({"status": "success"})


# 删除题目API
@app.route("/training_questions/delete/<int:qid>", methods=["POST"])
@login_required
def training_question_delete(qid):
    q = TrainingQuestion.query.filter_by(id=qid, user_id=current_user.id).first()
    if not q:
        return jsonify({"status": "not_found"})
    db.session.delete(q)
    db.session.commit()
    return jsonify({"status": "success"})


# ===============================
# 7. 培训系统 —— 培训任务管理/分配/答题/统计
# ===============================
# 新建培训任务及参与员工API
@app.route("/training_task/new", methods=["POST"])
@login_required
def create_training_task():
    title = request.form.get("title")
    material_id = request.form.get("material_id")
    description = request.form.get("description", "")
    deadline = request.form.get("deadline")
    employee_ids = request.form.getlist("employee_ids")
    max_attempts = int(request.form.get("max_attempts", 1))
    pass_score_ratio = float(request.form.get("pass_score_ratio", 80)) / 100.0

    if not (title and material_id and employee_ids):
        return jsonify({"status": "fail", "msg": "缺少必要参数"})

    # 校验该材料下题目数量
    num_questions = TrainingQuestion.query.filter_by(material_id=material_id).count()
    if num_questions == 0:
        return jsonify(
            {
                "status": "fail",
                "msg": "该培训材料没有任何题目，请先添加题库再发布培训任务。",
            }
        )

    # 1. 创建 TrainingTask
    task = TrainingTask(
        user_id=current_user.id,
        title=title,
        material_id=material_id,
        description=description,
        deadline=deadline if deadline else None,
        created_at=datetime.now(),
        max_attempts=max_attempts,
        pass_score_ratio=pass_score_ratio,
    )
    db.session.add(task)
    db.session.commit()

    # 2. 为每个员工生成 TrainingTaskEmployee 记录
    for emp_id in employee_ids:
        rec = TrainingTaskEmployee(task_id=task.id, employee_id=emp_id, status="未完成")
        db.session.add(rec)
    db.session.commit()
    return jsonify(
        {
            "status": "success",
            "redirect": url_for("training_task_invite", task_id=task.id),
        }
    )


# 培训任务详情页（员工及成绩等）
@app.route("/training_task/<int:task_id>")
@login_required
def training_task_detail(task_id):
    task = TrainingTask.query.filter_by(id=task_id, user_id=current_user.id).first()
    if not task:
        return "未找到该培训任务", 404

    # 查询培训材料
    material = TrainingMaterial.query.filter_by(id=task.material_id).first()

    # 查询参与员工及完成情况
    employees = (
        db.session.query(Employee, TrainingTaskEmployee)
        .join(TrainingTaskEmployee, Employee.id == TrainingTaskEmployee.employee_id)
        .filter(TrainingTaskEmployee.task_id == task_id)
        .all()
    )
    emp_list = [{"name": emp.name, "status": tte.status} for emp, tte in employees]

    # 生成培训答题链接（假设员工访问 /training_answer/<task_id> 开始答题）
    base_url = request.url_root.rstrip("/")
    answer_url = f"{base_url}/training_answer/{task_id}"

    return render_template(
        "training_task_invite.html",
        task=task,
        material=material,
        employees=emp_list,
        answer_url=answer_url,
    )


# 培训答题邀请页
@app.route("/training_task/invite/<int:task_id>")
@login_required
def training_task_invite(task_id):
    task = TrainingTask.query.get(task_id)
    material = TrainingMaterial.query.get(task.material_id)
    records = TrainingTaskEmployee.query.filter_by(task_id=task_id).all()
    emp_ids = [rec.employee_id for rec in records]
    employees = Employee.query.filter(Employee.id.in_(emp_ids)).all()
    # 查询每人历史
    histories = {}
    for emp in employees:
        hists = (
            TrainingAnswerHistory.query.filter_by(task_id=task.id, employee_id=emp.id)
            .order_by(TrainingAnswerHistory.attempt_num.asc())
            .all()
        )
        histories[emp.id] = [
            {
                "score": h.score,
                "is_passed": h.is_passed,
                "submit_time": h.submit_time.strftime("%Y-%m-%d %H:%M:%S"),
            }
            for h in hists
        ]

    emp_list = []
    for emp in employees:
        status = next(
            (rec.status for rec in records if rec.employee_id == emp.id), "未完成"
        )
        emp_list.append({"id": emp.id, "name": emp.name, "status": status})

    base_url = request.url_root.rstrip("/")
    return render_template(
        "training_task_invite.html",
        task=task,
        material=material,
        base_url=base_url,
        employees=emp_list,
        histories=histories,
    )


# 员工答题主入口，支持提交、计分、历史记录
@app.route("/training_answer/<int:task_id>/<int:employee_id>", methods=["GET", "POST"])
def training_answer(task_id, employee_id):
    # 1. 员工必须属于当前用户
    employee = Employee.query.filter_by(id=employee_id, user_id=current_user.id).first()
    if not employee:
        return (
            render_template(
                "training_answer_disabled.html", msg="员工不存在或无权限。"
            ),
            403,
        )

    # 2. 任务必须属于当前用户
    task = TrainingTask.query.filter_by(id=task_id, user_id=current_user.id).first()
    if not task:
        return (
            render_template(
                "training_answer_disabled.html", msg="培训任务不存在或无权限。"
            ),
            403,
        )

    # 3. 材料也要隔离
    material = TrainingMaterial.query.filter_by(
        id=task.material_id, user_id=current_user.id
    ).first()
    if not material:
        return (
            render_template(
                "training_answer_disabled.html", msg="培训材料不存在或无权限。"
            ),
            403,
        )

    # 4. 参与记录查找
    tte = TrainingTaskEmployee.query.filter_by(
        task_id=task_id, employee_id=employee_id
    ).first()
    if not tte:
        return (
            render_template(
                "training_answer_disabled.html", msg="未找到你的参与记录。"
            ),
            404,
        )

    # 5. 获取题库
    questions = TrainingQuestion.query.filter_by(
        material_id=material.id, user_id=current_user.id
    ).all()
    parsed_questions = [
        {
            "id": q.id,
            "content": q.content,
            "options": json.loads(q.options),
            "correct_answers": json.loads(q.correct_answers),
        }
        for q in questions
    ]

    # 6. 截止校验
    if task.deadline and datetime.now().date() > task.deadline:
        return render_template(
            "training_answer_disabled.html", msg="本次培训任务已截止，无法再答题。"
        )

    # 7. 是否已完成
    if tte.status == "已完成":
        return render_template(
            "training_answer_disabled.html",
            msg="你已完成本次答题，无需重复提交。",
            score=tte.score,
            is_passed=tte.is_passed,
            attempts=tte.attempts,
        )

    if request.method == "POST":
        if getattr(tte, "is_passed", False):
            return jsonify(
                {
                    "success": False,
                    "msg": "已通过，无需重复答题",
                    "passed": True,
                    "attempts": tte.attempts,
                    "max_attempts": task.max_attempts,
                }
            )
        if getattr(tte, "attempts", 0) >= task.max_attempts:
            tte.status = "已完成"
            db.session.commit()
            return jsonify(
                {
                    "success": False,
                    "msg": "已用完最大答题次数，无法再提交",
                    "passed": False,
                    "attempts": tte.attempts,
                    "max_attempts": task.max_attempts,
                }
            )

        answers = request.get_json()
        score = 0
        for q in parsed_questions:
            qid = f"q{q['id']}"
            ans = int(answers.get(qid, -1))
            if ans in q["correct_answers"]:
                score += 1

        tte.attempts = (tte.attempts or 0) + 1
        total = len(parsed_questions)
        correct_ratio = score / total if total else 0
        pass_score = getattr(task, "pass_score_ratio", 0.8)
        passed = correct_ratio >= pass_score

        history = TrainingAnswerHistory(
            task_id=task.id,
            employee_id=employee_id,
            attempt_num=tte.attempts,
            score=score,
            is_passed=passed,
            submit_time=datetime.now(),
        )
        db.session.add(history)

        tte.score = score
        tte.submit_time = datetime.now()
        tte.is_passed = passed

        if passed or tte.attempts >= task.max_attempts:
            tte.status = "已完成"
        else:
            tte.status = "未完成"

        db.session.commit()

        if not passed and tte.attempts >= task.max_attempts:
            return jsonify(
                {
                    "success": False,
                    "passed": False,
                    "score": score,
                    "total": total,
                    "msg": "未通过且已用完所有答题机会",
                    "attempts": tte.attempts,
                    "max_attempts": task.max_attempts,
                }
            )
        elif not passed:
            return jsonify(
                {
                    "success": False,
                    "passed": False,
                    "score": score,
                    "total": total,
                    "msg": f"未通过，可重试（剩余{task.max_attempts - tte.attempts}次）",
                    "attempts": tte.attempts,
                    "max_attempts": task.max_attempts,
                }
            )
        else:
            return jsonify(
                {
                    "success": True,
                    "passed": True,
                    "score": score,
                    "total": total,
                    "msg": "答题通过！",
                    "attempts": tte.attempts,
                    "max_attempts": task.max_attempts,
                }
            )

    return render_template(
        "training_answer.html",
        task=task,
        material=material,
        questions=parsed_questions,
        employee=employee,
        tte=tte,
        chr=chr,
    )


# 答题前选择员工身份页
@app.route("/training_answer_select/<int:task_id>", methods=["GET", "POST"])
def training_answer_select(task_id):
    # 只允许当前用户自己的任务
    task = TrainingTask.query.filter_by(id=task_id, user_id=current_user.id).first()
    if not task:
        return "未找到该培训任务或无权限", 404

    # 只找自己员工
    records = TrainingTaskEmployee.query.filter_by(task_id=task_id).all()
    emp_ids = [rec.employee_id for rec in records]

    # 员工必须属于当前用户
    employees = Employee.query.filter(
        Employee.id.in_(emp_ids), Employee.user_id == current_user.id
    ).all()

    if request.method == "POST":
        employee_id = int(request.form.get("employee_id"))
        # 仅允许选自己的员工
        if employee_id not in [e.id for e in employees]:
            return "无效的员工选择", 403
        return redirect(
            url_for("training_answer", task_id=task_id, employee_id=employee_id)
        )

    return render_template(
        "training_answer_select.html", task=task, employees=employees
    )


# 培训任务全局统计（供前端展示列表）
@app.route("/training_stats")
@login_required
def training_stats():
    # 只查当前用户自己的任务
    tasks = (
        TrainingTask.query.filter_by(user_id=current_user.id)
        .order_by(TrainingTask.created_at.desc())
        .all()
    )

    data = []
    for task in tasks:
        material = TrainingMaterial.query.get(task.material_id)
        ttes = TrainingTaskEmployee.query.filter_by(task_id=task.id).all()
        total = len(ttes)
        done = sum(1 for t in ttes if t.status == "已完成")
        data.append(
            {
                "id": task.id,
                "title": task.title,
                "material": material.title if material else "(无)",
                "created_at": task.created_at.strftime("%Y-%m-%d %H:%M"),
                "deadline": (
                    task.deadline.strftime("%Y-%m-%d") if task.deadline else "-"
                ),
                "progress": f"{done}/{total}",
                "is_completed": done == total and total > 0,
            }
        )
    return jsonify({"tasks": data})


# 删除培训任务API，级联删除所有关联
@app.route("/delete_training_task/<int:task_id>", methods=["POST"])
@login_required
def delete_training_task(task_id):
    # 只查当前用户自己的任务
    task = TrainingTask.query.filter_by(id=task_id, user_id=current_user.id).first()
    if not task:
        return jsonify({"status": "not_found"}), 404

    # 删除所有答题历史
    TrainingAnswerHistory.query.filter_by(task_id=task_id).delete()
    # 删除参与记录
    TrainingTaskEmployee.query.filter_by(task_id=task_id).delete()
    # 删除任务本身
    db.session.delete(task)
    db.session.commit()
    return jsonify({"status": "success"})


# 获取培训任务详情及统计数据API
@app.route("/training_task/get/<int:task_id>")
@login_required
def training_task_get(task_id):
    # 1. 只获取当前用户自己的任务
    task = TrainingTask.query.filter_by(id=task_id, user_id=current_user.id).first()
    if not task:
        return jsonify({"status": "not_found"})

    # 2. 获取培训材料（也做用户隔离更严谨，但最少做 task 隔离即可）
    material = TrainingMaterial.query.get(task.material_id)
    # 3. 只查本任务的参与员工
    ttes = TrainingTaskEmployee.query.filter_by(task_id=task_id).all()
    # 4. 批量获取员工对象
    employees = Employee.query.filter(
        Employee.id.in_([t.employee_id for t in ttes])
    ).all()
    emp_map = {e.id: e for e in employees}

    stats = {
        "total": len(ttes),
        "done": sum(1 for t in ttes if t.status == "已完成"),
        "avg_score": (
            round(sum((t.score or 0) for t in ttes) / len(ttes), 2) if ttes else 0
        ),
        "pass_ratio": (
            round(100.0 * sum(1 for t in ttes if t.is_passed) / len(ttes), 2)
            if ttes
            else 0
        ),
    }

    return jsonify(
        {
            "status": "success",
            "task": {
                "id": task.id,
                "title": task.title,
                "created_at": task.created_at.strftime("%Y-%m-%d %H:%M"),
                "material": material.title if material else "(无)",
                "deadline": task.deadline.strftime("%Y-%m-%d") if task.deadline else "",
                "description": task.description or "",
                "employees": [
                    {"id": t.employee_id, "name": emp_map[t.employee_id].name}
                    for t in ttes
                ],
                "stats": stats,
            },
        }
    )


# 导出培训记录PDF（Word→PDF）
@app.route("/api/export_training_record", methods=["POST"])
def export_training_record():
    # 获取表单前端发来的数据
    data = request.get_json()
    # 定位Word模板
    template_path = os.path.join(
        app.root_path, "static", "training_records", "培训记录表.docx"
    )
    doc = Document(template_path)

    # 填充Word表格中的各字段
    table = doc.tables[0]

    try:
        table.cell(1, 1).text = data.get("station", "")
        table.cell(2, 3).text = data.get("title", "")
        table.cell(2, 6).text = data.get("place", "")
        table.cell(3, 3).text = data.get("time", "")
        table.cell(3, 6).text = data.get("trainer", "")
        table.cell(4, 3).text = data.get("employees", "")

        # === 人数填充，run级别，不破坏格式 ===
        try:
            count = data.get("total")
            if not count:
                employees_str = data.get("employees", "")
                names = [
                    name.strip()
                    for name in re.split(r"[、，,\s\n]+", employees_str)
                    if name.strip()
                ]
                count = len(names)
            else:
                count = int(count)

            cell = table.cell(5, 2)
            para = cell.paragraphs[0]
            runs = para.runs
            found = False

            for run in runs:
                if run.text.strip() == "":
                    run.text = str(count)
                    found = True
                    break
            # 如果没有空run，则尝试用正则替换
            if not found:
                for run in runs:
                    if "共" in run.text and "人" in run.text:
                        run.text = re.sub(r"共\s*人", f"共 {count} 人", run.text)
                        found = True
                        break
            # 如果还是不行，最后兜底（用cell.text，样式会丢失，但保证内容）
            if not found:
                txt = cell.text
                cell.text = re.sub(r"共\s*人", f"共 {count} 人", txt)
        except Exception as e:
            print("人数填充异常", e)

        table.cell(6, 3).text = data.get("summary", "")
        table.cell(7, 3).text = data.get("result", "")
        table.cell(8, 3).text = data.get("note", "")
        table.cell(9, 3).text = data.get("date", "")
        table.cell(9, 6).text = data.get("reviewDate", "")
    except Exception as e:
        print("填表异常", e)
        return jsonify({"status": "fail", "msg": f"填表异常: {e}"}), 500

    # 保存为临时docx文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
        doc.save(tmp_docx.name)
        tmp_docx_path = tmp_docx.name

    # 转换为PDF，兼容不同操作系统
    pdf_path = tmp_docx_path.replace(".docx", ".pdf")
    sysplat = platform.system()
    try:
        if sysplat == "Linux":
            # 用libreoffice转换
            ret = os.system(
                f'libreoffice --headless --convert-to pdf "{tmp_docx_path}" --outdir "{os.path.dirname(tmp_docx_path)}"'
            )
            if ret != 0 or not os.path.exists(pdf_path):
                return (
                    jsonify(
                        {"status": "fail", "msg": "PDF转换失败，libreoffice执行错误"}
                    ),
                    500,
                )
        elif sysplat in ["Darwin", "Windows"]:
            try:
                from docx2pdf import convert

                convert(tmp_docx_path, pdf_path)
            except Exception as e:
                return jsonify({"status": "fail", "msg": f"docx2pdf转换失败: {e}"}), 500
        else:
            return (
                jsonify({"status": "fail", "msg": f"不支持的系统类型: {sysplat}"}),
                500,
            )
    except Exception as e:
        return jsonify({"status": "fail", "msg": f"PDF转换失败: {e}"}), 500

    # 返回PDF供前端下载
    return send_file(
        pdf_path,
        as_attachment=True,
        download_name=(data["title"] or "培训记录") + "-培训记录.pdf",
    )


# 保存培训记录表（持久化到数据库，覆盖同task_id旧数据）
@app.route("/api/save_training_record", methods=["POST"])
def save_training_record():
    # 1. 获取表单数据
    data = request.get_json()
    task_id = int(data.get("task_id"))
    # 2. 查询数据库有无历史记录，无则新建
    record = TrainingRecord.query.filter_by(
        task_id=task_id, user_id=current_user.id
    ).first()
    if not record:
        record = TrainingRecord(task_id=task_id, user_id=current_user.id)
        db.session.add(record)
    # 3. 更新所有字段
    record.station = data.get("station", "")
    record.title = data.get("title", "")
    record.time = data.get("time", "")
    record.place = data.get("place", "")
    record.trainer = data.get("trainer", "")
    record.employees = data.get("employees", "")
    record.summary = data.get("summary", "")
    record.result = data.get("result", "")
    record.note = data.get("note", "")
    record.date = data.get("date", "")
    record.review_date = data.get("review_date", "")
    record.updated_at = datetime.utcnow()
    try:
        db.session.commit()
        return jsonify({"status": "success"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"status": "fail", "msg": str(e)})


# 获取某培训任务的记录表（用于前端弹窗自动回填）
@app.route("/api/get_training_record/<int:task_id>", methods=["GET"])
def get_training_record(task_id):
    # 1. 查询数据库
    record = TrainingRecord.query.filter_by(
        task_id=task_id, user_id=current_user.id
    ).first()
    if not record:
        return jsonify({"status": "not_found"})
    # 2. 返回所有字段
    return jsonify(
        {
            "status": "success",
            "record": {
                "station": record.station or "",
                "title": record.title or "",
                "time": record.time or "",
                "place": record.place or "",
                "trainer": record.trainer or "",
                "employees": record.employees or "",
                "summary": record.summary or "",
                "result": record.result or "",
                "note": record.note or "",
                "date": record.date or "",
                "review_date": record.review_date or "",
            },
        }
    )


@app.route("/export_exam_docx", methods=["POST"])
def export_exam_docx():
    data = request.get_json()
    doc = Document()
    # 抬头
    p = doc.add_paragraph()
    r = p.add_run(data.get("header") or "（试卷抬头）")
    r.font.size = Pt(18)
    r.font.name = "方正小标宋简"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "方正小标宋简")
    p.alignment = 1
    # 标题
    p = doc.add_paragraph()
    r = p.add_run(data.get("title") or "（试卷标题）")
    r.font.size = Pt(22)
    r.font.name = "华文行楷"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "华文行楷")
    p.alignment = 1
    # 副标题
    p = doc.add_paragraph()
    r = p.add_run(data.get("subtitle") or "（副标题）")
    r.font.size = Pt(18)
    r.font.name = "华文楷体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "华文楷体")
    p.alignment = 1
    # 考试信息
    p = doc.add_paragraph()
    p.alignment = 1
    r = p.add_run(
        f"考试时长：{data['time']}    满分：{data['score']}    及格线：{data['pass']}"
    )
    r.font.size = Pt(12)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    # 姓名栏
    p = doc.add_paragraph()
    p.alignment = 1
    r = p.add_run(
        "姓名：__________    得分：__________    站名：__________    岗位：__________"
    )
    r.font.size = Pt(12)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.add_paragraph("")  # 空行

    # 题目分组
    judge = [
        q
        for q in data["questions"]
        if q.get("qtype") == "judge"
        or (not q.get("options") or len(q.get("options", [])) == 2)
    ]
    single = [
        q
        for q in data["questions"]
        if q.get("qtype") == "single"
        or (q.get("options") and len(q.get("options", [])) > 2)
    ]

    judge_total = 40
    single_total = 60
    judge_score = round(judge_total / len(judge), 2) if judge else 0
    single_score = round(single_total / len(single), 2) if single else 0

    # 判断题
    if judge:
        p = doc.add_paragraph()
        run = p.add_run(f"一、判断题（每题{judge_score}分，共{judge_total}分）")
        run.font.size = Pt(12)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        for idx, q in enumerate(judge):
            p = doc.add_paragraph()
            # 括号加在题号前
            r = p.add_run(f"（   ）{idx+1}. {q.get('content')}")
            r.font.size = Pt(12)
            r.font.name = "宋体"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 单选题
    if single:
        p = doc.add_paragraph()
        run = p.add_run(f"二、单项选择题（每题{single_score}分，共{single_total}分）")
        run.font.size = Pt(12)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        for idx, q in enumerate(single):
            p = doc.add_paragraph()
            r = p.add_run(f"{idx+1}. {q.get('content')}（   ）")  # 括号加在题目最后
            r.font.size = Pt(12)
            r.font.name = "宋体"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            option_lines = smart_layout_options(q.get("options", []))  # 智能排布选项
            for line in option_lines:
                po = doc.add_paragraph("    " + line)
                runo = po.runs[0]
                runo.font.size = Pt(12)
                runo.font.name = "宋体"
                runo._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.add_paragraph("")
    p = doc.add_paragraph()
    r = p.add_run("阅卷人签字：_____________________        日期：______________")
    r.font.size = Pt(12)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 保存临时Word文件并返回
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_word:
        doc.save(tmp_word.name)
        word_path = tmp_word.name

    return send_file(
        word_path, as_attachment=True, download_name=(data["title"] or "试卷") + ".docx"
    )


def get_char_len(s):
    # 简单估算字符串显示宽度：中文算2，英文数字算1
    return sum(2 if ord(c) > 127 else 1 for c in s)


def smart_layout_options(options, max_line_len=36):
    """根据内容长度智能排列选项，每行不超过max_line_len字符"""
    if not options:
        return []
    n = len(options)
    abcd = [f"{chr(65+i)}、{opt}" for i, opt in enumerate(options)]
    # 1. 尝试一行放下
    joined = "    ".join(abcd)
    if get_char_len(joined) <= max_line_len:
        return [joined]
    # 2. 尝试两行平均分
    if n == 4:
        first = "    ".join(abcd[:2])
        second = "    ".join(abcd[2:])
        if max(get_char_len(first), get_char_len(second)) <= max_line_len:
            return [first, second]
    if n == 3:
        first = "    ".join(abcd[:2])
        second = abcd[2]
        if max(get_char_len(first), get_char_len(second)) <= max_line_len:
            return [first, second]
    # 3. 每行一个
    return abcd


# ===============================
# 8. 启动/运维/其它扩展
# ===============================
if __name__ == "__main__":
    with app.app_context():
        db.create_all()  # 首次启动创建所有表
    app.run(host="0.0.0.0", port=5050, debug=True)
